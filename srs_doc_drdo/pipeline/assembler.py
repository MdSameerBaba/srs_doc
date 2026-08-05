"""
assembler.py — Stitches individual SRS sections into a final IEEE 830 document.
Also generates a verification summary report.
"""

import json
from datetime import datetime


import io


def create_docx_bytes(md_text: str, project_name: str = "SRS Document") -> bytes:
    """
    Convert Markdown SRS text into a Microsoft Word (.docx) file bytes stream.
    Supports python-docx if installed, with a pure-Python OpenXML zip fallback.
    """
    try:
        import docx
        from docx.shared import Inches, Pt
        doc = docx.Document()
        
        # Style Title
        title_p = doc.add_heading(f"Software Requirements Specification — {project_name}", level=0)
        title_p.paragraph_format.space_after = Pt(12)
        
        for line in md_text.splitlines():
            line_s = line.strip()
            if not line_s:
                continue
            if line_s.startswith("# "):
                doc.add_heading(line_s[2:], level=1)
            elif line_s.startswith("## "):
                doc.add_heading(line_s[3:], level=2)
            elif line_s.startswith("### "):
                doc.add_heading(line_s[4:], level=3)
            elif line_s.startswith("#### "):
                doc.add_heading(line_s[5:], level=4)
            elif line_s.startswith("- ") or line_s.startswith("* "):
                doc.add_paragraph(line_s[2:], style='List Bullet')
            elif line_s.startswith("> "):
                p = doc.add_paragraph(line_s[2:])
                p.paragraph_format.left_indent = Inches(0.5)
            elif line_s.startswith("---"):
                continue
            else:
                doc.add_paragraph(line_s)
                
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        pass

    # Pure-Python OpenXML (.docx) zip builder fallback
    import zipfile
    import xml.sax.saxutils as saxutils

    def escape_xml(s: str) -> str:
        return saxutils.escape(s)

    p_xml_list = []
    for line in md_text.splitlines():
        line_s = line.strip()
        if not line_s or line_s.startswith("---"):
            continue
        text_esc = escape_xml(line_s)
        if line_s.startswith("# "):
            p_xml_list.append(f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="36"/></w:rPr><w:t>{text_esc[2:]}</w:t></w:r></w:p>')
        elif line_s.startswith("## "):
            p_xml_list.append(f'<w:p><w:pPr><w:pStyle w:val="Heading2"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>{text_esc[3:]}</w:t></w:r></w:p>')
        elif line_s.startswith("### "):
            p_xml_list.append(f'<w:p><w:pPr><w:pStyle w:val="Heading3"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="24"/></w:rPr><w:t>{text_esc[4:]}</w:t></w:r></w:p>')
        elif line_s.startswith("- ") or line_s.startswith("* "):
            p_xml_list.append(f'<w:p><w:pPr><w:pStyle w:val="ListBullet"/></w:pPr><w:r><w:t>• {text_esc[2:]}</w:t></w:r></w:p>')
        else:
            p_xml_list.append(f'<w:p><w:r><w:t>{text_esc}</w:t></w:r></w:p>')

    doc_body = "".join(p_xml_list)

    content_types_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>'
    rels_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    document_xml = f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>{doc_body}</w:body></w:document>'

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("word/document.xml", document_xml)

    return buf.getvalue()


# ─────────────────────────────────────────────────────────────
# Final SRS assembly
# ─────────────────────────────────────────────────────────────

def assemble_srs(
    sections: dict,
    canonical: dict,
    project_name: str = "Unknown Project",
) -> str:
    """
    Combine all generated sections into a single Markdown SRS document.
    Section 11 (Traceability Matrix) is auto-generated from canonical FRs
    if not already present in `sections`.
    """
    norm_sections: dict[int, str] = {}
    for k, v in sections.items():
        try:
            if isinstance(k, str):
                num = int(k.split("_")[0])
            else:
                num = int(k)
            norm_sections[num] = str(v)
        except Exception:
            pass

    ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    header = f"""\
# Software Requirements Specification

**Project:** {project_name}
**Standard:** IEEE 830 / ISO/IEC 29148
**Generated:** {ts}
**Generator:** AI SRS Generator (Pure-Python AST + Ollama)

> ⚠️ This document was auto-generated from source code analysis.
> All `[TBD]` entries require manual completion before formal release.
> Verify every `[evidence: ...]` citation against the frozen `canonical.json`.

---

"""

    body_parts: list[str] = [header]

    for section_num in sorted(norm_sections.keys()):
        md = norm_sections[section_num].strip()
        body_parts.append(md)
        body_parts.append("\n\n---\n\n")

    # ── Auto-generate Section 11 if missing ────────────────
    if 11 not in norm_sections:
        body_parts.append(_build_traceability_matrix(canonical))
        body_parts.append("\n\n---\n\n")

    return "".join(body_parts)


def _build_traceability_matrix(canonical: dict) -> str:
    from pathlib import Path
    frs = canonical.get("functional_requirements", []) or []
    nfrs = canonical.get("non_functional_signals", []) or []
    
    # ── 11.1 Requirements to Source File Traceability ────
    rows_11_1 = [
        "## 11. Traceability Matrix\n",
        "### 11.1 Requirements to Source File Traceability\n",
        "| REQ-ID | Requirement Summary | Section | Source File(s) |",
        "|---|---|---|---|",
    ]
    
    traceable_count = 0
    derived_count = 0
    
    # Process functional requirements
    for fr in frs:
        fr_id = fr.get("id", "")
        title = (fr.get("title") or "")[:70].replace("|", "\\|")
        ev_ids = fr.get("evidence_ids") or []
        
        files_set = set()
        for ev in ev_ids:
            if ":" in ev:
                parts = ev.split(":")
                if len(parts) >= 2 and "." in parts[1]:
                    files_set.add(parts[1])
        
        if files_set:
            ev_str = ", ".join(sorted(files_set))
            traceable_count += 1
        else:
            ev_str = "Derived from system design"
            derived_count += 1
            
        rows_11_1.append(f"| {fr_id} | {title} | Section 7 | {ev_str} |")
        
    # Process non-functional requirements
    for idx, nfr in enumerate(nfrs):
        nfr_id = f"REQ-NF-{idx+1:03d}"
        desc = (nfr.get("description") or "")[:70].replace("|", "\\|")
        ev_ids = nfr.get("evidence_ids") or []
        
        files_set = set()
        for ev in ev_ids:
            if ":" in ev:
                parts = ev.split(":")
                if len(parts) >= 2 and "." in parts[1]:
                    files_set.add(parts[1])
                    
        if files_set:
            ev_str = ", ".join(sorted(files_set))
            traceable_count += 1
        else:
            ev_str = "Derived from system design"
            derived_count += 1
            
        rows_11_1.append(f"| {nfr_id} | {desc} | Section 7 | {ev_str} |")

    # ── 11.2 Requirements to System Feature Traceability ────
    rows_11_2 = [
        "\n### 11.2 Requirements to System Feature Traceability\n",
        "| REQ-ID | Requirement Summary | System Feature / Module | Verification Method |",
        "|---|---|---|---|",
    ]
    
    for fr in frs:
        fr_id = fr.get("id", "")
        title = (fr.get("title") or "")[:70].replace("|", "\\|")
        ev_ids = fr.get("evidence_ids") or []
        
        module_name = "Core System"
        for ev in ev_ids:
            if ":" in ev:
                parts = ev.split(":")
                if len(parts) >= 2 and "." in parts[1]:
                    module_name = Path(parts[1]).stem
                    break
        
        title_lower = title.lower()
        if any(w in title_lower for w in ["ui", "display", "screen", "render", "show", "view"]):
            method = "Demonstration"
        elif any(w in title_lower for w in ["performance", "speed", "fast", "time", "latency", "scale"]):
            method = "Analysis"
        else:
            method = "Test"
            
        rows_11_2.append(f"| {fr_id} | {title} | {module_name} | {method} |")
        
    for idx, nfr in enumerate(nfrs):
        nfr_id = f"REQ-NF-{idx+1:03d}"
        desc = (nfr.get("description") or "")[:70].replace("|", "\\|")
        ev_ids = nfr.get("evidence_ids") or []
        
        module_name = "Core System"
        for ev in ev_ids:
            if ":" in ev:
                parts = ev.split(":")
                if len(parts) >= 2 and "." in parts[1]:
                    module_name = Path(parts[1]).stem
                    break
                    
        cat = nfr.get("category", "").lower()
        if "perf" in cat or "scal" in cat:
            method = "Analysis"
        elif "sec" in cat or "rel" in cat:
            method = "Analysis"
        else:
            method = "Review"
            
        rows_11_2.append(f"| {nfr_id} | {desc} | {module_name} | {method} |")

    # ── 11.3 Traceability Summary ────
    total_fr = len(frs)
    total_nfr = len(nfrs)
    
    summary_text = f"""
### 11.3 Traceability Summary
- Total number of functional requirements identified: **{total_fr}**
- Total number of non-functional requirements identified: **{total_nfr}**
- Number of requirements traceable to specific source files: **{traceable_count}**
- Number of requirements derived from system design (no direct source file evidence): **{derived_count}**
"""
    
    return "\n".join(rows_11_1) + "\n" + "\n".join(rows_11_2) + "\n" + summary_text


# ─────────────────────────────────────────────────────────────
# Verification summary report
# ─────────────────────────────────────────────────────────────

def generate_verification_report(verification_results: dict[int, dict]) -> str:
    """Return a Markdown report summarising per-section audit results."""
    pass_count = sum(
        1 for r in verification_results.values()
        if isinstance(r, dict) and r.get("status") == "PASS"
    )
    total = len(verification_results)

    lines = [
        "# Verification Report",
        f"\n**{pass_count}/{total} sections passed verification**\n",
    ]

    for section_num in sorted(verification_results.keys()):
        report = verification_results[section_num]
        if not isinstance(report, dict):
            continue
        status = report.get("status", "UNKNOWN")
        icon   = "✅" if status == "PASS" else ("⚠️" if status == "NEEDS_REVIEW" else "❌")
        lines.append(f"### {icon} Section {section_num} — `{status}`")

        if report.get("missing"):
            lines.append(f"- **Missing FRs:** {', '.join(report['missing'])}")
        if report.get("hallucinated"):
            lines.append(f"- **Hallucinated FRs:** {', '.join(report['hallucinated'])}")
        if report.get("renamed"):
            for r in report["renamed"]:
                lines.append(
                    f"- **Renamed:** `{r.get('id')}` "
                    f"canonical=*{r.get('canonical_title')}* "
                    f"→ generated as *{r.get('generated_as')}*"
                )
        if report.get("uncited_claims"):
            lines.append(
                f"- **Uncited claims:** {len(report['uncited_claims'])} "
                f"(first: *{report['uncited_claims'][0][:80]}*)"
            )
        lines.append("")

    return "\n".join(lines)
